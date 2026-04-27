# -*- coding: utf-8 -*-
"""
DeviceBrief サンプル記事 → Word (.docx) 変換
sample_articles.md を読み込んで DeviceBrief_Sample_Articles.docx を生成
"""

import sys
import io
import re
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_MD   = Path(__file__).parent / "sample_articles.md"
OUTPUT_DOC = Path(__file__).parent / "DeviceBrief_Sample_Articles_v3.docx"

# ─────────────────────────────────────────────
# ヘルパー
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """表セルの背景色を設定"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_horizontal_rule(doc: Document):
    """区切り線を追加"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 40)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(9)

def add_meta_table(doc: Document, meta: dict):
    """メタデータを2列テーブルで表示"""
    if not meta:
        return
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, val in meta.items():
        row = table.add_row()
        cell_k = row.cells[0]
        cell_v = row.cells[1]
        cell_k.width = Cm(4)
        set_cell_bg(cell_k, "E8F0FE")
        p_k = cell_k.paragraphs[0]
        run_k = p_k.add_run(key)
        run_k.bold = True
        run_k.font.size = Pt(8)
        run_k.font.color.rgb = RGBColor(0x23, 0x4A, 0x9E)
        p_v = cell_v.paragraphs[0]
        run_v = p_v.add_run(str(val))
        run_v.font.size = Pt(8)
    doc.add_paragraph()

# ─────────────────────────────────────────────
# Markdownパーサー（sample_articles.md用）
# ─────────────────────────────────────────────

def parse_articles(md_text: str) -> list[dict]:
    """
    ## 記事N: ... で区切られた記事ブロックを解析
    戻り値: [{header, meta, body_lines}]
    """
    # 記事ブロックを分割（## 記事N: で始まる行）
    blocks = re.split(r"(?=^## 記事\d+:)", md_text, flags=re.MULTILINE)
    articles = []

    for block in blocks:
        block = block.strip()
        if not block.startswith("## 記事"):
            continue

        # ヘッダー行（## 記事N: ...）
        header_match = re.match(r"^## (記事\d+: .+)", block)
        header = header_match.group(1) if header_match else ""

        # ---META--- ブロック抽出
        meta = {}
        meta_match = re.search(r"---META---(.*?)---END_META---", block, re.DOTALL)
        if meta_match:
            for line in meta_match.group(1).strip().splitlines():
                if ":" in line:
                    k, _, v = line.partition(":")
                    meta[k.strip()] = v.strip()

        # METAブロックを除いた本文
        body = re.sub(r"---META---.*?---END_META---\n?", "", block, flags=re.DOTALL)
        # ヘッダー行を除去
        body = re.sub(r"^## 記事\d+:.*\n?", "", body).strip()

        articles.append({
            "header": header,
            "meta": meta,
            "body": body,
        })

    return articles


def render_body(doc: Document, body: str):
    """Markdown本文をdocxに変換"""
    lines = body.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # H2見出し
        if line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=2)
            h.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
            i += 1
            continue

        # 引用ブロック（信頼度★）
        if line.startswith("> "):
            p = doc.add_paragraph(style="Quote")
            run = p.add_run(line[2:])
            run.font.size = Pt(10)
            run.bold = True
            run.font.color.rgb = RGBColor(0xD9, 0x73, 0x06)
            i += 1
            continue

        # 出典セクション
        if line.startswith("**出典**"):
            p = doc.add_paragraph()
            run = p.add_run("出典")
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            i += 1
            # 出典リスト
            while i < len(lines) and lines[i].startswith("- "):
                src_line = lines[i][2:]
                # [タイトル](URL) — メディア のパース
                m = re.match(r"\[(.+?)\]\((.+?)\)\s*(?:—|-)\s*(.+)", src_line)
                p2 = doc.add_paragraph(style="List Bullet")
                p2.paragraph_format.left_indent = Cm(0.5)
                if m:
                    run2 = p2.add_run(f"{m.group(1)}")
                    run2.font.size = Pt(8)
                    run2.font.color.rgb = RGBColor(0x23, 0x4A, 0x9E)
                    run2.underline = True
                    p2.add_run(f"  — {m.group(3)}").font.size = Pt(8)
                else:
                    p2.add_run(src_line).font.size = Pt(8)
                i += 1
            continue

        # 区切り線
        if line.strip() == "---":
            i += 1
            continue

        # 空行
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # 通常テキスト（**bold** パース）
        p = doc.add_paragraph()
        parts = re.split(r"(\*\*[^*]+\*\*)", line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)
            run.font.size = Pt(10.5)
        i += 1

# ─────────────────────────────────────────────
# Word生成メイン
# ─────────────────────────────────────────────

def build_docx():
    md_text = INPUT_MD.read_text(encoding="utf-8")
    articles = parse_articles(md_text)
    print(f"パース完了: {len(articles)} 記事")

    doc = Document()

    # ── ページ余白設定 ──
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── 表紙 ──
    title_p = doc.add_heading("DeviceBrief", 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("サンプル記事集 — 品質チェック用")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    from datetime import datetime
    date_run = date_p.add_run(f"生成日: {datetime.now().strftime('%Y年%m月%d日')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    doc.add_paragraph()
    doc.add_paragraph()

    # 記事一覧
    index_h = doc.add_heading("収録記事一覧", level=1)
    index_h.runs[0].font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    for a in articles:
        art_type = a["meta"].get("article_type", "")
        category = a["meta"].get("category", "")
        title_val = a["meta"].get("title", a["header"])
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(f"[{art_type} / {category}] {title_val}")
        run.font.size = Pt(10)

    doc.add_page_break()

    # ── 各記事 ──
    for idx, art in enumerate(articles):
        # 記事番号ヘッダー
        header_p = doc.add_heading(art["header"], level=1)
        header_p.runs[0].font.color.rgb = RGBColor(0x11, 0x18, 0x27)

        # メタデータ表
        if art["meta"]:
            meta_disp = {k: v for k, v in art["meta"].items()
                         if k in ("title", "article_type", "category", "tags",
                                  "source_reliability", "seo_description")}
            add_meta_table(doc, meta_disp)

        # 本文
        render_body(doc, art["body"])

        # 次の記事との区切り（最後の記事は不要）
        if idx < len(articles) - 1:
            doc.add_page_break()

    doc.save(OUTPUT_DOC)
    print(f"保存完了: {OUTPUT_DOC}")
    print(f"ファイルサイズ: {OUTPUT_DOC.stat().st_size:,} bytes")


if __name__ == "__main__":
    build_docx()
