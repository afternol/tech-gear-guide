# -*- coding: utf-8 -*-
"""DeviceBrief concept.docx generator - UTF-8"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)

def tbl(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    h_row = t.rows[0]
    for i, h in enumerate(headers):
        c = h_row.cells[i]
        c.text = h
        for run in c.paragraphs[0].runs:
            run.bold = True
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9E1F2")
        c._tc.get_or_add_tcPr().append(shd)
    for row_data in rows:
        r = t.add_row()
        for i, text in enumerate(row_data):
            r.cells[i].text = text

def bl(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

doc = Document()
doc.styles["Normal"].font.name = "Meiryo UI"
doc.styles["Normal"].font.size = Pt(10.5)

# ---- Title ----
h = doc.add_heading("DeviceBriefプロジェクト全方针・決定事項", level=0)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("作成日: 2026-04-24　|　ステータス: 方针確定・開発準備中")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
hr(doc)

# ---- 1 ----
doc.add_heading("1. メディア基本情報", level=1)
tbl(doc, ["項目", "内容"], [
    ["サイト名", "DeviceBrief"],
    ["ドメイン", "devicebrief.com 等（取得予定・未確定）"],
    ["コンセプト", "海外権威メディアの情報を正確に伝え、専門家視点の示唑を加える日本語テックメディア"],
    ["ターゲット読者", "テックリテラシー高めの日本人（20～40代）"],
    ["競合参考サイト", "tech-gadget.reinforz.co.jp（記事フォーマットのみ参考・他は独自）"],
    ["運営", "Reinforz Insight"],
])
doc.add_paragraph()

# ---- 2 ----
doc.add_heading("2. 取り扱いジャンル・カテゴリ", level=1)
bl(doc, [
    "ニュース（速報）",
    "iPhone / iOS",
    "Android",
    "Windows",
    "CPU・GPU・チップ",
    "アプリ・ソフトウェア",
    "周辺機器・ガジェット",
])
doc.add_paragraph()

# ---- 3 ----
doc.add_heading("3. 記事方针・品質ルール", level=1)
doc.add_heading("3-1. 記事タイプ（3種）", level=2)
tbl(doc, ["タイプ", "文字数", "内容", "1日の本数目標"], [
    ["A型・速報", "1,500～2,000字", "海夰1～2ソースの要点＋専門解説", "35～40本"],
    ["B型・深掴り", "1,500～2,000字", "複数ソース統合・業界背景・日本市場への影響", "5～8本"],
    ["C型・リーク正説", "1,500～2,000字", "次世代機の歬説＋ソース信頼度評価＋確度表示", "5～8本"],
])
p = doc.add_paragraph()
p.add_run("1日合計目標: 50記事").bold = True

doc.add_heading("3-2. 記事構成（固定ルール）", level=2)
bl(doc, [
    "見出し数: H2を必ず4つ",
    "文字数: 1,500～2,000字（全タイプ共通）",
    "出典: 必ずURL付きで記載（最低2本）",
    "リード文: 凒頭に2～3行（見出しなし）",
])

doc.add_heading("3-3. 品質ルール（厳守）", level=2)
p = doc.add_paragraph()
r = p.add_run("「必須」")
r.bold = True
r.font.color.rgb = RGBColor(0, 112, 192)
bl(doc, [
    "ソース記事の事実を正確に伝える（数字・スペック・日付はソース通り）",
    "ソース信頼度の明示（Bloombergが報道、リーカー情報）",
    "専門的示唑を含める（なぜ重要か・業界へのインパクト・日本ユーザーへの意味）",
    "出典URLを全記事に最低2本必ず記載",
])
p = doc.add_paragraph()
r = p.add_run("「禁止」")
r.bold = True
r.font.color.rgb = RGBColor(192, 0, 0)
bl(doc, [
    "ソースにない情報の捧造（ハルシネーション絶対禁止）",
    "推測をファクトとして記述",
    "出典なしの記事投稿",
    "翻訳だけの薄い記事（示唑・解説が必須）",
])

doc.add_heading("3-4. C型リーク記事の信頼度表示", level=2)
tbl(doc, ["確度", "対象"], [
    ["★★★★★", "Mark Gurman（Bloomberg）、Zac Bowden（Windows Central）"],
    ["★★★★☆", "9to5Mac / 9to5Google 独自情報"],
    ["★★★☆☆", "著名リーカー（@OnLeaks、Ice Universe等）"],
    ["★★☆☆☆", "匿名ソース・未確認情報"],
    ["★☆☆☆☆", "Reddit・SNS上の未確認リーク"],
])
doc.add_paragraph()

# ---- 4 ----
doc.add_heading("4. ニュースソース設計", level=1)

doc.add_heading("4-1. スマホ系", level=2)
tbl(doc, ["Tier", "ソース", "強み"], [
    ["1", "The Verge", "総合テック最高権威"],
    ["1", "9to5Mac", "iPhone速報の決定版"],
    ["1", "9to5Google", "Pixel/Android最速"],
    ["1", "Android Authority", "Samsung/Xiaomi幅広い"],
    ["1", "GSMArena", "スペック・レビュー"],
    ["1", "Bloomberg Technology", "Gurman情報・産業動向"],
    ["2", "MacRumors", "iPhone歪説・リーク"],
    ["2", "PhoneArena", "比較レビュー"],
    ["2", "XDA Developers", "技術深掴り"],
    ["2", "TechRadar", "総合ガジェット"],
    ["2", "Engadget", "トレンド・実用"],
    ["2", "NotebookCheck", "リーク最速・ベンチマーク"],
    ["3", "Reddit r/Androidヽr/iphoneヽr/GooglePixel", "バズ検知"],
])
doc.add_paragraph()

doc.add_heading("4-2. Windows系", level=2)
tbl(doc, ["Tier", "ソース", "強み"], [
    ["1", "Windows Central", "Windows専門No.1"],
    ["1", "Ars Technica", "技術深掴り・信頼性高"],
    ["2", "Neowin", "Windows速報"],
    ["2", "Thurrott", "Microsoftエコシステム"],
    ["3", "Reddit r/Windows11ヽr/Surface", "バズ検知"],
])
doc.add_paragraph()

doc.add_heading("4-3. CPU・GPU・チップ系", level=2)
tbl(doc, ["Tier", "ソース", "強み"], [
    ["1", "Tom's Hardware", "CPU/GPU最重要メディア"],
    ["1", "Ars Technica", "チップ・半導体深掴り"],
    ["2", "Wccftech", "GPU/CPUリーク"],
    ["2", "VideoCardz", "GPU専門"],
    ["2", "NotebookCheck", "ラップトップ・チップ性能"],
    ["3", "Reddit r/hardwareヽr/AMDヽr/nvidia", "バズ検知"],
])
doc.add_paragraph()

doc.add_heading("4-4. 産業・ビジネス系", level=2)
tbl(doc, ["Tier", "ソース", "強み"], [
    ["1", "Reuters Technology", "ファクト重視の速報"],
    ["1", "Financial Times Tech", "産業・企業動向"],
    ["2", "The Information", "内部情報・調査報道"],
])
doc.add_paragraph()

doc.add_heading("4-5. 著名リーカー監視（SNS）", level=2)
bl(doc, [
    "Mark Gurman (@markgurman) — Apple最高信頼度",
    "Ice Universe (@UniverseIce) — Samsung",
    "@OnLeaks (Steve Hemmerstoffer) — 3Dレンダリング",
    "Yogesh Brar (@heyitsyogesh) — Android情報",
    "WalkingCat (@_h0x0d_) — Microsoft内部リーク",
    "Zac Bowden (Windows Central) — Microsoft公式筋",
])
doc.add_paragraph()

# ---- 5 ----
doc.add_heading("5. 著者方针", level=1)
p = doc.add_paragraph()
p.add_run("方针: 「DeviceBrief編集部」方式（架空ライター不使用）").bold = True
doc.add_paragraph("Googleの現行方针ではAIコンテンツ自体は問題ないが、AIを人間と偽ることはペナルティリスク。架空ライターは2024年以降ペナルティ増加傾向。編集部方式＋透明性強化でE-E-A-Tを確保する。")
doc.add_heading("E-E-A-T強化策", level=2)
bl(doc, [
    "Aboutページ: 海兤20社以上を常時モニタリングする編集方针を明記",
    "編集方针ページ: 情報精度最優先・ソース確認プロセス・誤報訂正フローを記載",
    "AIアシスト明示: 各記事にAI支援により作成の一行（信頼性シグナル）",
    "運営会社の透明性: Reinforz Insightとの関係を明示",
    "出典の徹底: 全記事に出典URLを必ず記載",
])
doc.add_paragraph()

# ---- 6 ----
doc.add_heading("6. 技術スタック", level=1)
tbl(doc, ["レイヤー", "技術", "理由"], [
    ["フロントエンド", "Next.js 16 (App Router)", "SEO・ISR対応・既存プロジェクトと統一"],
    ["UI", "shadcn/ui", "既存プロジェクトと統一"],
    ["データベース", "Supabase（Postgres）", "無料枚で十分・成長に耐えられる"],
    ["画像ストレージ", "Supabase Storage", "記事アイキャッチ画像の保管"],
    ["ホスティング", "Vercel", "Next.jsとの相性最良・ISR対応"],
    ["CMS", "なし（パイプライン直接投稿）", "WordPressは使用しない"],
    ["パイプライン", "Python（SmaTechから改造）", "既存資産の流用"],
    ["自動化", "GitHub Actions（Cron）", "1日3回自動実行"],
])
doc.add_paragraph()

# ---- 7 ----
doc.add_heading("7. Supabase データベース設計（Phase 1）", level=1)
doc.add_paragraph("articles テーブルの主要カラム:")
bl(doc, [
    "id (uuid, PK)",
    "title (text)",
    "slug (text, unique)",
    "body (text) — Markdown本文",
    "category (text)",
    "tags (text[])",
    "article_type (text) — A型/B型/C型",
    "source_reliability (int) — C型のみ 1～5",
    "sources (jsonb) — [{title, url, media}]",
    "featured_image_url (text)",
    "seo_description (text)",
    "published_at (timestamptz)",
    "created_at (timestamptz)",
    "is_published (bool)",
])
doc.add_paragraph()

# ---- 8 ----
doc.add_heading("8. コンテンツ生成パイプライン", level=1)
doc.add_heading("SmaTechからの差分", level=2)
tbl(doc, ["項目", "SmaTech", "DeviceBrief"], [
    ["sources.yaml", "スマホのみ", "スマホ+Windows+CPU/GPU追加"],
    ["generate.py", "翻訳+要約中心", "専門示唑+信頼度評価+H2x4固定"],
    ["publish.py", "WordPress REST API", "Supabase REST API（INSERT）"],
    ["記事タイプ", "A型・B型", "A型・B型・C型（リーク）"],
    ["文字数", "800～3,000字（可変）", "1,500～2,000字（全タイプ統一）"],
    ["見出し数", "2～4個（可変）", "H2x4個（固定）"],
    ["画像処理", "WPメディアアップロード", "Supabase Storage"],
    ["自動化", "手動実行", "GitHub Actions Cron（1日3回）"],
])
doc.add_paragraph()

doc.add_heading("Claude APIコスト試算", level=2)
bl(doc, [
    "1記事: 約 2,500 output tokens",
    "50記事/日: 125,000 output tokens",
    "費用: $0.38/日 = ¥55/日 = ¥1,650/月",
    "（入力トークン含め ¥3,000～4,000/月 想定）",
])
doc.add_paragraph()

# ---- 9 ----
doc.add_heading("9. SEO戦略", level=1)
doc.add_heading("技術SEO", level=2)
bl(doc, [
    "Core Web Vitals 最適化（LCP < 2.5秒）",
    "構造化データ（NewsArticle / BreadcrumbList / FAQPage）",
    "XMLサイトマップ自動生成",
    "canonical URL",
    "Google News 申請・登録",
    "robots.txt 適切設定",
])
doc.add_heading("コンテンツSEO", level=2)
bl(doc, [
    "タイトル: デバイス名+モデル番号+年号（例: iPhone 17 Pro スペック判明《2026年》）",
    "メタ description: 150字以内で要点凝縮",
    "H2x4の論理的構造",
    "内部リンク（同一デバイスシリーズ記事を自動リンク）",
    "タイトルパターン: 「～が判明」「～を確認」「～の可能性」「～とは？」",
])
doc.add_paragraph()

# ---- 10 ----
doc.add_heading("10. AIO（AI検索最適化）戦略", level=1)
doc.add_paragraph("ChatGPT / Perplexity / Google AI Overviews に引用されるための設計:")
bl(doc, [
    "事実の密度を高める — 数字・日付・スペック・価格を明確に記述",
    "問いへの直接回答構造 — 検索意図に即答するリード文",
    "FAQセクション — 各記事末尾に3～5問のQ&A（将来実装）",
    "出典の透明性 — Bloomberg の Mark Gurman が報告のような出典明記",
    "E-E-A-T シグナル — 編集方针・About・Privacy Policy の充実",
    "エンティティ最適化 — デバイス名・モデル番号・メーカー名を正確に記述",
])
doc.add_paragraph()

# ---- 11 ----
doc.add_heading("11. Google Discovery戦略", level=1)
bl(doc, [
    "アイキャッチ画像: 1,200x628px以上（高品質）",
    "タイトル: 好奇心を刺激する（「～が判明」「～の衝撃スペック」「～が激変」）",
    "新鮮さ: 公開後30分以内にインデックス（ISR活用）",
    "Google News 登録済みサイト",
    "話題のデバイスは即日記事化（iPhone次世代・Galaxy新作等）",
    "サムネイルに実機画像を使用（OGP取得 or 公式プレスイメージ）",
    "1日50本の更新頻度維持（Discovery配信量に直結）",
])
doc.add_paragraph()

# ---- 12 ----
doc.add_heading("12. マネタイズ戦略（3フェーズ）", level=1)
tbl(doc, ["フェーズ", "期間", "PV目標", "収益源"], [
    ["Phase 1\n記事蒂積期", "0～3ケ月", "月隙10万PV", "Google AdSense（10万PV到達で申請）"],
    ["Phase 2\nSEO流入期", "3～6ケ月", "月隙50万PV", "AdSense+Amazonアフィリエイト+A8.net\n目標: 月¥3～10万"],
    ["Phase 3\n収益最大化期", "6ケ月～", "月隙200万PV", "プレスリリース港載（¥30,000～/本）\n記事スポンサーシップ\n有料ニュースレター\n売却・M&A"],
])
doc.add_paragraph()

# ---- 13 ----
doc.add_heading("13. DB×成長戦略", level=1)
tbl(doc, ["フェーズ", "DB追加機能", "目的"], [
    ["Phase 1", "articlesテーブルのみ", "記事保存・配信"],
    ["Phase 2", "page_viewsテーブル\npopular_articlesビュー\nrelated_articles", "人気記事分析\n内部リンク自動生成"],
    ["Phase 3", "user_sessionsテーブル\nkeywordsテーブル\ntrending_topics", "行動分析\nSearch Console連携\nトレンド自動検出"],
])
doc.add_paragraph()

# ---- 14 ----
doc.add_heading("14. コンテンツ鮮度維持・流入減衰対策", level=1)
doc.add_paragraph("速報記事100%構成では、デバイス世代交代時に関連記事が一括陳腐化し流入が急落する。以下9つの対策を全て実施する。")

doc.add_heading("コンテンツポートフォリオ設計", level=2)
tbl(doc, ["タイプ", "比率", "寿命", "内容"], [
    ["速報ニュース", "70%", "1～7日", "Discovery流入の主力"],
    ["デバイスまとめ（継続更新）", "15%", "恒久", "ハブ記事・随時更新"],
    ["エバーグリーン", "15%", "1～3年", "「使い方」「選び方」「比較」"],
])
doc.add_paragraph("エバーグリーンは速報50本/日とは別に週5～10本追加生成する。")

doc.add_heading("主要対策一覧", level=2)
bl(doc, [
    "ハブ&スポーク構造: デバイスまとめハブ記事に全速報をリンク。新スポーク公開時にupdated_atを自動更新。",
    "記事自動リフレッシュ: 前週比50%以下の記事を検知し、週次バッチで自動再生成。同一URLを維持。",
    "デバイスライフサイクル対応: 発表前→速報→深掘り→まとめの自動シフト。",
    "年間カレンダー: CES/MWC/WWDC/iPhone発表等を2週間前から仕込み。",
    "メールマガジン: beehiiv or Resendで週次配信。アルゴリズム非依存の直接流入確保。",
    "プラットフォーム分散: X自動投稿、Google News登録、RSS/Atom配信。",
    "内部リンク自動管理: 新記事公開時に関連旧記事へ自動追記。PageRankを維持。",
    "流入減衰の自動検知: Supabase + Search Console API連携でリフレッシュキューを自動管理。",
])
doc.add_paragraph()

# ---- 15 ----
doc.add_heading("15. 圧倒的ユーザー価値・SEO1位・Discovery最大化戦略", level=1)
doc.add_paragraph("競合の翻訳+要約サイトに対し、以下9戦略で物理的に真似できない価値を設計する。")

doc.add_heading("戦略1: 速報パイプライン革新（15〜30分公開）", level=2)
doc.add_paragraph("常駐Pythonプロセス（5分ごとRSSポーリング）でPriority Queueを運用。CRITICAL（発表・リーク）は15分以内、HIGH（スペック・価格）は30分以内に公開。GitHub ActionsはNORMAL記事専用に降格。実装: VPS or Railway.app。")

doc.add_heading("戦略2: Progressive Article（他サイトがやっていない最大の差別化）", level=2)
tbl(doc, ["フェーズ", "タイミング", "文字数", "内容"], [
    ["Phase 1 速報", "T+15分", "600～800字", "事実のみ。LIVEバッジ表示。"],
    ["Phase 2 詳細", "T+2～6時間", "1,500～2,000字", "複数ソース統合。日本市場影響。技術詳細。"],
    ["Phase 3 決定版", "T+24～48時間", "2,000～2,500字", "全情報統合。比較表。FAQ。今後の予測。"],
])
doc.add_paragraph("同一URL・同一スラッグで成長する。SEO/Discovery/AIO全てで最強の構造。")

doc.add_heading("戦略3: 日本向け独自価値（唯一無二の情報）", level=2)
tbl(doc, ["独自要素", "内容", "実装"], [
    ["JPY自動換算", "価格は必ず円換算（為替API連携）", "Python"],
    ["日本発売日", "日本での発売日・入手可否を必ず調査", "プロンプト指示"],
    ["キャリア対応", "ドコモ・au・ソフトバンク・楽天対応状況", "プロンプト指示"],
    ["技適・周波数", "日本バンド対応（SIMフリー購入判断）", "プロンプト指示"],
    ["日本版の差異", "グローバル版と日本版の仕様差", "プロンプト指示"],
])
doc.add_paragraph()

doc.add_heading("戦略4: 記事の7層品質構造", level=2)
bl(doc, [
    "レイヤー1: 事実の正確な報告（基本）",
    "レイヤー2: なぜ重要か（業界文脈）",
    "レイヤー3: 日本市場への影響（独自価値）",
    "レイヤー4: 技術的深掘り（専門性）",
    "レイヤー5: 前世代・競合との比較（判断支援）→スペックDB連携で自動生成",
    "レイヤー6: 今後の予測（付加価値）",
    "レイヤー7: FAQ（AIO対策 + 検索意図カバレッジ）",
])

doc.add_heading("戦略5: OGPサムネイル自動生成", level=2)
doc.add_paragraph("Pythonで記事ごとに1,200x628px画像を動的生成。デバイス実機画像＋カテゴリバッジ（LEAK/速報/独自）＋タイトル文字＋ロゴ。Discovery上のCTRが大幅向上。")

doc.add_heading("戦略6: SEO1位設計", level=2)
bl(doc, [
    "Top Storiesカルーセル: Google News登録 + NewsArticle構造化データ（dateModified必須）。Progressive Articleの更新のたびに再配信対象になる。",
    "Featured Snippets（位置0）: ハウツー記事は番号付きリスト使用。最初の200字で直接回答。",
    "People Also Ask（PAA）: 全記事末尾にFAQセクション3～5問 + FAQPageスキーマ実装。PAAボックスを独占。",
])

doc.add_heading("戦略7: タイトル設計公式", level=2)
bl(doc, [
    "速報系: 「[デバイス名] [事実]が[判明/確定/流出]【[ソース名]報道】」",
    "リーク系: 「[デバイス名]のレンダリング画像が流出—[衝撃的変更点]」",
    "日本向け: 「[デバイス名]の日本価格・発売日が判明—[円換算]から」",
    "エバーグリーン: 「[デバイス名]の[機能名]を最大限活かす[N]の方法【2026年版】」",
])

doc.add_heading("戦略8: ユーザー定着設計", level=2)
tbl(doc, ["機能", "効果", "実装"], [
    ["PWAプッシュ通知", "速報をリアルタイム配信→直接流入", "Next.js PWA"],
    ["デバイス追跡機能", "特定デバイスの最新情報をフォロー", "Supabase"],
    ["週次メルマガ", "アルゴリズム非依存の直接流入", "beehiiv or Resend"],
    ["カテゴリ別RSSフィード", "テック系ユーザーの購読", "Next.js API Route"],
])
doc.add_paragraph()

doc.add_heading("戦略9: データ蓄積による参入障壁", level=2)
tbl(doc, ["テーブル", "内容", "活用記事例"], [
    ["benchmarks", "GeekbenchスコアをDB化", "「歴代チップベンチマーク比較」"],
    ["specs", "全デバイスのスペックをDB化", "「スペック比較ページ」"],
    ["prices", "価格推移をDB化", "「値下がり予測・買い時判断」"],
])
doc.add_paragraph("競合が後から追いつこうとしても蓄積の時間差がある。参入障壁になる。")

doc.add_heading("全体方程式", level=2)
doc.add_paragraph(
    "速報性（15分以内）x Progressive Article（記事が成長する）x 日本独自価値（円換算・発売日・キャリア）"
    "x 7層品質構造（比較表・信頼度・FAQ・予測）x SEO設計（Top Stories・Featured Snippet・PAA）"
    "x Discovery最適化（OGP動的生成・タイトル公式）x ユーザー定着（PWA・メルマガ・RSS）"
    "x データ蓄積（スペックDB・ベンチマークDB）x 鮮度維持（ハブ記事・エバーグリーン・リフレッシュ）"
    " = 日本最高のテックデバイスメディア"
)
doc.add_paragraph()

# ---- 16 ----
doc.add_heading("16. 未決定事項（要確認）", level=1)
bl(doc, [
    "ドメイン取得（devicebrief.com 等を調査・確保）",
    "Vercel or Xserver ホスティング選定",
    "Supabaseプロジェクト新規作成 or 既存流用",
    "GitHub Actions の実行スケジュール最終確定（7時・12時・18時 案）",
    "Google News 申請タイミング（記事50本以上が目安）",
    "AdSense 申請タイミング（10万PV到達後）",
])

output_path = r"C:\Users\after\devicebrief\docs\devicebrief_concept.docx"
doc.save(output_path)
print("OK: " + output_path)
